from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_base_template():
    doc = Document()
    
    # 1. KOP SURAT
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("{{ nama_sekolah|default('NAMA SEKOLAH ANDA DI SINI', true) }}\n")
    run.bold = True
    run.font.size = Pt(16)
    
    run2 = p.add_run("TAHUN AJARAN 2026/2027\n")
    run2.font.size = Pt(11)
    
    run3 = p.add_run("Jalan Contoh Nomor 123, Kota Anda, Telp: (021) 123456\n")
    run3.font.size = Pt(10)
    
    # Pemisah garis ganda tidak bisa mudah di docx murni via script, jadi user bisa menambahkannya nanti.
    doc.add_paragraph("_" * 70)
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("LEMBAR SOAL\n")
    run.bold = True
    run.font.size = Pt(12)
    
    # 2. METADATA
    table = doc.add_table(rows=4, cols=3)
    data = [
        ("Mata Pelajaran", ":", "{{ mata_pelajaran }}"),
        ("Kelas", ":", "{{ kelas }}"),
        ("Hari/Tanggal", ":", "{{ tanggal }}"),
        ("Waktu", ":", "{{ durasi }}")
    ]
    for i, row in enumerate(table.rows):
        row.cells[0].text = data[i][0]
        row.cells[1].text = data[i][1]
        row.cells[2].text = data[i][2]
        # Atur lebar sedikit
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(0.2)
        row.cells[2].width = Inches(3.0)
        
    doc.add_paragraph("\nPETUNJUK UMUM")
    doc.add_paragraph("1. Bacalah Basmalah sebelum memulai mengerjakan soal.\n2. Tulislah nama dan nomor peserta pada lembar jawaban.\n3. Periksa dan bacalah soal dengan teliti sebelum menjawab.")
    
    doc.add_paragraph("_" * 70)
    
    # 3. BLOK SOAL (menggunakan sintaks docxtpl)
    
    # Pilihan Ganda
    doc.add_paragraph("{% if pg_list|length > 0 %}")
    p = doc.add_paragraph("I. PILIHLAH SALAH SATU JAWABAN YANG PALING TEPAT!")
    p.runs[0].bold = True
    
    doc.add_paragraph("{% for s in pg_list %}")
    doc.add_paragraph("{{ s.nomor }}. {{ s.teks }}")
    doc.add_paragraph("{% for o in s.opsi %}")
    doc.add_paragraph("    {{ o.label }}. {{ o.teks }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endif %}\n")
    
    # Isian
    doc.add_paragraph("{% if isian_list|length > 0 %}")
    p = doc.add_paragraph("II. ISILAH TITIK-TITIK DI BAWAH INI DENGAN TEPAT!")
    p.runs[0].bold = True
    
    doc.add_paragraph("{% for s in isian_list %}")
    doc.add_paragraph("{{ s.nomor }}. {{ s.teks }}")
    doc.add_paragraph("    Jawaban: ..............................................................\n")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endif %}\n")
    
    # Essay
    doc.add_paragraph("{% if essay_list|length > 0 %}")
    p = doc.add_paragraph("III. JAWABLAH PERTANYAAN DI BAWAH INI DENGAN BENAR!")
    p.runs[0].bold = True
    
    doc.add_paragraph("{% for s in essay_list %}")
    doc.add_paragraph("{{ s.nomor }}. {{ s.teks }}\n\n\n")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endif %}")

    # Buat direktori template jika belum ada
    os.makedirs('template', exist_ok=True)
    doc.save('template/template_ujian.docx')
    print("Template berhasil dibuat di template/template_ujian.docx")

if __name__ == '__main__':
    create_base_template()
