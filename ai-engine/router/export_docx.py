import logging
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
import io

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    pass

router = APIRouter()
logger = logging.getLogger(__name__)

class HeaderDocx(BaseModel):
    nama_sekolah: str = Field(default="")
    mata_pelajaran: str = Field(default="")
    kelas: str = Field(default="")
    tanggal: str = Field(default="")
    durasi: str = Field(default="")

class RequestExportDocx(BaseModel):
    soal: list = Field(min_length=1)
    header: HeaderDocx
    sertakan_jawaban: bool = Field(default=False)
    sertakan_pembahasan: bool = Field(default=False)

@router.post("/export/docx")
async def export_docx(req: RequestExportDocx):
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx belum ter-install. Jalankan: pip install python-docx",
        )
        
    try:
        doc = Document()
        
        # Header
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(req.header.nama_sekolah or "Soal Ujian\n")
        run.bold = True
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info = f"Mata Pelajaran: {req.header.mata_pelajaran} | Kelas: {req.header.kelas}\n"
        info += f"Tanggal: {req.header.tanggal} | Durasi: {req.header.durasi}"
        p.add_run(info)
        
        doc.add_paragraph("-" * 80)
        
        # Soal
        nomor = 1
        for s in req.soal:
            tipe = s.get("tipe", s.get("type", "pg"))
            teks = s.get("teks", s.get("teks_soal", s.get("text", "")))
            
            doc.add_paragraph(f"{nomor}. {teks}")
            
            if tipe == "pg":
                opsi_list = s.get("opsi", s.get("options", []))
                for opsi in opsi_list:
                    label = opsi.get("label", opsi.get("label_opsi", ""))
                    teks_opsi = opsi.get("teks", opsi.get("teks_opsi", opsi.get("text", "")))
                    p_opsi = doc.add_paragraph(f"   {label}. {teks_opsi}")
            elif tipe == "essay":
                doc.add_paragraph("\n\n\n")
            elif tipe == "isian":
                doc.add_paragraph("   Jawaban: _________________________________")
                
            doc.add_paragraph("") # spacing
            nomor += 1
            
        # Jawaban
        if req.sertakan_jawaban:
            doc.add_page_break()
            doc.add_heading('Kunci Jawaban', level=1)
            for i, s in enumerate(req.soal, 1):
                kunci = s.get("kunci_jawaban", s.get("answer_key", "-"))
                doc.add_paragraph(f"{i}. {kunci}")
                
        # Pembahasan
        if req.sertakan_pembahasan:
            doc.add_page_break()
            doc.add_heading('Pembahasan', level=1)
            for i, s in enumerate(req.soal, 1):
                pembahasan = s.get("pembahasan", s.get("explanation", "-"))
                doc.add_paragraph(f"{i}. {pembahasan}")
                
        # Save to memory
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return StreamingResponse(
            f,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=soal-educraft.docx"},
        )
        
    except Exception as e:
        logger.error(f"Export DOCX gagal: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
